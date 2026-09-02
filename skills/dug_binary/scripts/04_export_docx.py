"""
DUG inventory xlsx → per-volume process DOCX (step 4 of 5).

Reads the step-3 workbook (Volume_Processes sheet, UUIDs already resolved)
and writes a Word document: two H1 sections (volumes with no derived
processes — alphabetical; volumes with derived processes — alphabetical),
one heading + process table per volume, per-volume Parents/Children lineage
mini-tables, and a call-out box pointing to the HTML lineage explorer
(step 5).

--no-pid produces the IDs-stripped variant: pid suffixes removed from
headings, pid column dropped, «name» (vol/pidN) → «name» (vol), portrait
layout retained. The script asserts the literal string "pid" appears
nowhere in the no-pid output.

Usage:
    python 04_export_docx.py --xlsx path/to/inventory.xlsx --out path/to/out.docx [--no-pid]
    python 04_export_docx.py --selftest
"""
import argparse, datetime as dt, re, sys
from collections import OrderedDict

PID_REF = re.compile(r"«([^»]+)» \(([^/()]+)/pid(\d+)\)")
BLUE, ORANGE, GREY = "1F4E78", "C55A11", "BFBFBF"


def strip_pids(text):
    """Remove every literal pid reference from a parameter string."""
    return PID_REF.sub(lambda m: f"«{m.group(1)}» ({m.group(2)})", text or "")


def load_rows(xlsx):
    from openpyxl import load_workbook
    wb = load_workbook(xlsx, read_only=True, data_only=True)
    ws = wb["Volume_Processes"]
    rows = list(ws.iter_rows(min_row=3, values_only=True))  # r1 banner, r2 headers
    wb.close()
    return rows


def group_volumes(rows):
    """rows: (vpid, vname, dpid, dtype, ptype, dname, params). Returns
    (with_procs OrderedDict[(vpid,vname)] -> [(dpid,ptype,dname,params)], without list)."""
    withp, without = OrderedDict(), []
    for vpid, vname, dpid, dtype, ptype, dname, params in rows:
        if dpid in (None, "—"):
            without.append((vpid, vname))
            continue
        withp.setdefault((vpid, vname), []).append((dpid, ptype, dname, params))
    without = sorted(set(without), key=lambda t: str(t[1] or "").lower())
    withp = OrderedDict(sorted(withp.items(), key=lambda kv: str(kv[0][1] or "").lower()))
    return withp, without


def build_lineage(withp):
    """Parents of V = vol-pid refs inside V's parameter strings (self-loops
    skipped). Children = inverse. Returns pid -> {'name', 'parents', 'children'}
    with (pid, name) tuples."""
    names = {vpid: vname for (vpid, vname) in withp}
    parents, children = {}, {}
    for (vpid, vname), procs in withp.items():
        for _dpid, _ptype, _dname, params in procs:
            for m in PID_REF.finditer(params or ""):
                if m.group(2) != "vol":
                    continue
                ppid = int(m.group(3))
                if ppid == vpid:
                    continue
                names.setdefault(ppid, m.group(1))
                parents.setdefault(vpid, set()).add(ppid)
                children.setdefault(ppid, set()).add(vpid)
    out = {}
    for pid in set(names) | set(parents) | set(children):
        out[pid] = {
            "name": names.get(pid, f"(pid {pid})"),
            "parents": sorted(parents.get(pid, ()), key=lambda p: str(names.get(p, "")).lower()),
            "children": sorted(children.get(pid, ()), key=lambda p: str(names.get(p, "")).lower()),
        }
    return out


# ---------------------------------------------------------------- docx build

def _borders(tbl):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4"); b.set(qn("w:color"), GREY)
        borders.append(b)
    tbl._tbl.tblPr.append(borders)


def _shade(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(sh)


def _header_row(tbl, labels, fill, widths=None):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for i, label in enumerate(labels):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(label)
        r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        _shade(cell, fill)
        if widths:
            cell.width = widths[i]


def _center_table(tbl):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER


def export_docx(xlsx, out, no_pid=False, explorer_name="Volume_Lineage_Explorer.html"):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rows = load_rows(xlsx)
    withp, without = group_volumes(rows)
    lineage = build_lineage(withp)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)  # portrait A4
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.5))
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    title = "Volume Processes" + (" (IDs removed)" if no_pid else "")
    doc.add_heading(title, level=0)
    p = doc.add_paragraph()
    p.add_run("Source: ").bold = True
    p.add_run(f"{xlsx} (sheet: Volume_Processes)")
    p = doc.add_paragraph()
    p.add_run("Generated: ").bold = True
    p.add_run(dt.datetime.now(dt.timezone.utc).strftime("%Y-%b-%d %H:%M UTC"))
    p = doc.add_paragraph()
    p.add_run(f"Volumes with derived processes: ").bold = True
    p.add_run(f"{len(withp)}; ")
    p.add_run("with no derived-process link: ").bold = True
    p.add_run(f"{len(without)}.")

    # Call-out box → interactive explorer (replaces embedded diagrams)
    box = doc.add_table(rows=1, cols=1)
    _borders(box); _center_table(box)
    cell = box.rows[0].cells[0]
    _shade(cell, "FFF2CC")
    cell.text = ""
    r = cell.paragraphs[0].add_run(
        "Volume lineage (parent–child) diagrams are provided as an interactive, "
        f"self-contained HTML explorer: {explorer_name}. Open it in any browser "
        "(offline, no install) — hover to enlarge, click to pin a volume and "
        "highlight its parents (blue) and children (orange), search, pan and zoom."
    )
    r.font.size = Pt(9)
    doc.add_paragraph()

    def volume_heading(vpid, vname):
        h = doc.add_heading(level=1)
        run = h.add_run(f"{vname or '(unnamed volume)'}")
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        if not no_pid:
            pr = h.add_run(f"  [vol pid {vpid}]")
            pr.font.size = Pt(11); pr.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    def lineage_tables(vpid):
        from docx.shared import Pt as _Pt
        meta = lineage.get(vpid)
        if not meta or not (meta["parents"] or meta["children"]):
            return
        for kind, fill in (("parents", BLUE), ("children", ORANGE)):
            pids = meta[kind]
            if not pids:
                continue
            tbl = doc.add_table(rows=1, cols=1)
            _borders(tbl); _center_table(tbl)
            label = kind.capitalize()
            _header_row(tbl, [f"{label} ({len(pids)})"], fill)
            for p_ in pids:
                row = tbl.add_row()
                nm = lineage[p_]["name"] if p_ in lineage else f"(pid {p_})"
                txt = nm if no_pid else f"{nm}  (pid {p_})"
                run = row.cells[0].paragraphs[0].add_run(txt)
                run.font.size = _Pt(9)

    # Section 1 — no derived processes (alphabetical)
    doc.add_heading(f"Volumes with no derived processes ({len(without)})", level=1)
    doc.add_paragraph(
        "Direct imports or virtual volumes whose lineage isn't captured via "
        "AttributeRevision.relatedProductId."
    ).runs[0].italic = True
    cols = 1 if no_pid else 2
    tbl = doc.add_table(rows=1, cols=cols)
    _borders(tbl); _center_table(tbl)
    _header_row(tbl, ["Volume name"] if no_pid else ["Volume pid", "Volume name"], BLUE)
    for vpid, vname in without:
        row = tbl.add_row()
        vals = [str(vname or "")] if no_pid else [str(vpid), str(vname or "")]
        for i, v in enumerate(vals):
            row.cells[i].text = v

    # Section 2 — with derived processes (alphabetical)
    doc.add_page_break()
    doc.add_heading(f"Volumes with derived processes ({len(withp)})", level=1)
    if no_pid:
        widths = [Cm(3.0), Cm(5.5), Cm(9.5)]
        labels = ["ProcessType", "Derived name / description", "Parameters (key=value; …)"]
    else:
        widths = [Cm(2.0), Cm(3.0), Cm(5.0), Cm(8.0)]
        labels = ["Derived pid", "ProcessType", "Derived name / description", "Parameters (key=value; …)"]

    for (vpid, vname), procs in withp.items():
        volume_heading(vpid, vname)
        p = doc.add_paragraph(); p.add_run(f"{len(procs)} derived process(es).").italic = True
        tbl = doc.add_table(rows=1, cols=len(labels))
        tbl.autofit = False
        _borders(tbl); _center_table(tbl)
        _header_row(tbl, labels, BLUE, widths)
        for dpid, ptype, dname, params in procs:
            params = strip_pids(params) if no_pid else (params or "")
            vals = ([str(ptype or ""), str(dname or ""), params] if no_pid
                    else [str(dpid), str(ptype or ""), str(dname or ""), params])
            row = tbl.add_row()
            for i, v in enumerate(vals):
                cell = row.cells[i]
                cell.text = ""
                run = cell.paragraphs[0].add_run(v)
                run.font.size = Pt(9)
                if labels[i].startswith("Parameters"):
                    run.font.name = "Consolas"
                cell.width = widths[i]
        lineage_tables(vpid)

    doc.save(out)

    if no_pid:
        # Hard guarantee: zero literal "pid" occurrences in the document XML text.
        from docx import Document as _D
        texts = []
        d = _D(out)
        for para in d.paragraphs:
            texts.append(para.text)
        for t in d.tables:
            for rw in t.rows:
                for c in rw.cells:
                    texts.append(c.text)
        joined = "\n".join(texts).lower()
        assert "pid" not in joined, "no-pid output still contains literal 'pid'"
    return len(withp), len(without)


def selftest():
    assert strip_pids("in=«VolA» (vol/pid12524); x=1") == "in=«VolA» (vol); x=1"
    rows = [
        (1, "B vol", 10, 17, "VolumeMaths", "d1", "in=«A vol» (vol/pid2)"),
        (2, "A vol", None, None, None, None, None),
        (3, "C vol", 11, 17, "Filter", "d2", "in=«C vol» (vol/pid3)"),  # self-loop
    ]
    withp, without = group_volumes(rows)
    assert [v for _, v in without] == ["A vol"]
    lin = build_lineage(withp)
    assert lin[1]["parents"] == [2] and lin[2]["children"] == [1]
    assert lin[3]["parents"] == [] and lin[3]["children"] == []  # self-loop skipped
    print("selftest OK")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--xlsx", help="step-3 inventory workbook")
    ap.add_argument("--out", help="output docx path")
    ap.add_argument("--no-pid", action="store_true", help="strip all product IDs")
    ap.add_argument("--explorer-name", default="Volume_Lineage_Explorer.html",
                    help="filename referenced in the call-out box")
    ap.add_argument("--selftest", action="store_true")
    a = ap.parse_args()
    if a.selftest:
        selftest(); sys.exit(0)
    if not (a.xlsx and a.out):
        ap.error("--xlsx and --out are required (no defaults — user must supply paths)")
    nw, nwo = export_docx(a.xlsx, a.out, no_pid=a.no_pid, explorer_name=a.explorer_name)
    print(f"OK -> {a.out}  ({nw} volumes with processes, {nwo} without)")
